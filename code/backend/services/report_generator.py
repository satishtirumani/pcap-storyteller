"""
PCAP StoryTeller - Forensic Report Generator
-------------------------------------------
This service creates professional PDF and Word documents summarizing 
the forensic investigation findings for students and stakeholders.
"""

import io
from flask import send_file, jsonify
from data.data_manager import DataManager
from services.services import AnalyticsService

def _prepare_data_for_reporting(events_list):
    """
    Cleans and formats raw dictionary data into readable table rows.
    
    Args:
        events_list (list): Raw list of forensic events from the parser.
        
    Returns:
        tuple: (list of formatted rows, dict of event type frequencies)
    """
    from datetime import datetime
    import json
    
    formatted_rows = []
    type_frequency_map = {}
    
    for event in events_list:
        event_category = event.get('type', 'Unknown')
        type_frequency_map[event_category] = type_frequency_map.get(event_category, 0) + 1
        
        # Step: Convert the computer-system timestamp into a human-readable date/time
        raw_timestamp = event.get('timestamp', 0)
        try: 
            readable_time = datetime.fromtimestamp(float(raw_timestamp)).strftime('%Y-%m-%d %H:%M:%S')
        except Exception: 
            readable_time = str(raw_timestamp)
        
        # Build a single readable row for the report tables
        formatted_rows.append({
            'time': readable_time,
            'type': event_category,
            'src': event.get('source_ip') or 'Internal',
            'dst': event.get('dest_ip') or 'Internal',
            'description': event.get('description') or 'No description provided',
            'details': json.dumps(event.get('details', {}), indent=2)
        })
        
    return formatted_rows, type_frequency_map

def generate_pdf_report():
    """
    Constructs a visual PDF document using the reportlab library.
    """
    # 1. Load the analyzed findings from the Data Manager
    forensic_data = DataManager.load_findings()
    if not forensic_data or not forensic_data.get('events'):
        return jsonify({'error': 'No analysis data found. Please parse a PCAP file first.'}), 400

    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    # 2. Prepare our forensic data into clean rows and counts
    table_rows, event_counts = _prepare_data_for_reporting(forensic_data['events'])
    
    # 3. Create a memory buffer to hold the document while we build it
    memory_buffer = io.BytesIO()
    pdf_doc = SimpleDocTemplate(memory_buffer, pagesize=letter, title='Security Forensic Report')
    text_styles = getSampleStyleSheet()
    document_elements = []

    # 4. Step: Add Title and Branding
    document_elements.append(Paragraph('🛡️ PCAP StoryTeller: Forensic Investigation Report', text_styles['Title']))
    document_elements.append(Spacer(1, 12))

    # 5. Step: Add High-Level Summary Table
    document_elements.append(Paragraph('Analysis Summary', text_styles['Heading2']))
    summary_table_data = [['Forensic Category', 'Incident Count']]
    for category, total in sorted(event_counts.items()):
        summary_table_data.append([category, str(total)])
    
    summary_table = Table(summary_table_data, colWidths=[300, 100])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
    ]))
    document_elements.append(summary_table)
    document_elements.append(Spacer(1, 24))

    # 6. Step: Add Detailed Activity Log (Limiting to top 100 for file size)
    document_elements.append(Paragraph('Detailed Forensic Activity (First 100 Events)', text_styles['Heading2']))
    detailed_table_headers = [['Time', 'Type', 'Source IP', 'Dest IP', 'Action Summary']]
    for row in table_rows[:100]:
        detailed_table_headers.append([
            row['time'], 
            row['type'], 
            row['src'], 
            row['dst'], 
            row['description'][:50] # Truncate for table layout
        ])

    log_table = Table(detailed_table_headers, colWidths=[90, 70, 90, 90, 150])
    log_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 8)
    ]))
    document_elements.append(log_table)

    # 7. Finalize and serve the PDF file to the browser
    pdf_doc.build(document_elements)
    memory_buffer.seek(0)
    return send_file(
        memory_buffer, 
        as_attachment=True, 
        download_name='pcap-analysis-report.pdf', 
        mimetype='application/pdf'
    )

def generate_docx_report():
    """
    Constructs a Microsoft Word (.docx) document using the python-docx library.
    """
    forensic_data = DataManager.load_findings()
    if not forensic_data: 
        return jsonify({'error': 'No data found.'}), 400

    from docx import Document
    table_rows, event_counts = _prepare_data_for_reporting(forensic_data['events'])
    
    # 1. Initialize the Word Document
    word_doc = Document()
    word_doc.add_heading('PCAP StoryTeller: Forensic Investigation Report', 0)
    
    # 2. Add Summary Table
    word_doc.add_heading('High-Level Summary', level=1)
    summary_word_table = word_doc.add_table(rows=1, cols=2)
    summary_word_table.style = 'Table Grid'
    header_cells = summary_word_table.rows[0].cells
    header_cells[0].text = 'Event Type'
    header_cells[1].text = 'Occurrence Count'
    
    for category, count in sorted(event_counts.items()):
        new_row = summary_word_table.add_row().cells
        new_row[0].text = category
        new_row[1].text = str(count)

    # 3. Add Detailed Log (Top 100)
    word_doc.add_heading('Activity Detail Log (Top 100)', level=1)
    detailed_word_table = word_doc.add_table(rows=1, cols=5)
    detailed_word_table.style = 'Table Grid'
    log_headers = detailed_word_table.rows[0].cells
    for i, title in enumerate(['Time', 'Type', 'Source', 'Dest', 'Description']):
        log_headers[i].text = title

    for row in table_rows[:100]:
        new_cells = detailed_word_table.add_row().cells
        new_cells[0].text = row['time']
        new_cells[1].text = row['type']
        new_cells[2].text = row['src']
        new_cells[3].text = row['dst']
        new_cells[4].text = row['description']

    # 4. Stream the file back to the user
    memory_buffer = io.BytesIO()
    word_doc.save(memory_buffer)
    memory_buffer.seek(0)
    return send_file(
        memory_buffer, 
        as_attachment=True, 
        download_name='pcap-analysis-report.docx', 
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
